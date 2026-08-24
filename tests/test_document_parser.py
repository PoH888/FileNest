import base64
from hashlib import sha256
from pathlib import Path
from uuid import UUID

import pytest
from docx import Document as WordDocument
from docx.enum.section import WD_SECTION

from backend.app.document_parser import (
    DocumentParseError,
    UnsupportedDocumentFormatError,
    load_document,
)
from backend.app.filesystem_adapter import FileSystemAdapter


DOCUMENT_ID = UUID("8a96e8c4-6ebc-4b5f-83b8-b27c681e8d95")
PDF_FIXTURE_BASE64 = (
    "JVBERi0xLjMKJZOMi54gUmVwb3J0TGFiIEdlbmVyYXRlZCBQREYgZG9jdW1lbnQgKG9wZW5zb3Vy"
    "Y2UpCjEgMCBvYmoKPDwKL0YxIDIgMCBSCj4+CmVuZG9iagoyIDAgb2JqCjw8Ci9CYXNlRm9udCAv"
    "SGVsdmV0aWNhIC9FbmNvZGluZyAvV2luQW5zaUVuY29kaW5nIC9OYW1lIC9GMSAvU3VidHlwZSAv"
    "VHlwZTEgL1R5cGUgL0ZvbnQKPj4KZW5kb2JqCjMgMCBvYmoKPDwKL0NvbnRlbnRzIDggMCBSIC9N"
    "ZWRpYUJveCBbIDAgMCAzMDAgMjAwIF0gL1BhcmVudCA3IDAgUiAvUmVzb3VyY2VzIDw8Ci9Gb250"
    "IDEgMCBSIC9Qcm9jU2V0IFsgL1BERiAvVGV4dCAvSW1hZ2VCIC9JbWFnZUMgL0ltYWdlSSBdCj4+"
    "IC9Sb3RhdGUgMCAvVHJhbnMgPDwKCj4+IAogIC9UeXBlIC9QYWdlCj4+CmVuZG9iago0IDAgb2Jq"
    "Cjw8Ci9Db250ZW50cyA5IDAgUiAvTWVkaWFCb3ggWyAwIDAgMzAwIDIwMCBdIC9QYXJlbnQgNyAw"
    "IFIgL1Jlc291cmNlcyA8PAovRm9udCAxIDAgUiAvUHJvY1NldCBbIC9QREYgL1RleHQgL0ltYWdl"
    "QiAvSW1hZ2VDIC9JbWFnZUkgXQo+PiAvUm90YXRlIDAgL1RyYW5zIDw8Cgo+PiAKICAvVHlwZSAv"
    "UGFnZQo+PgplbmRvYmoKNSAwIG9iago8PAovUGFnZU1vZGUgL1VzZU5vbmUgL1BhZ2VzIDcgMCBS"
    "IC9UeXBlIC9DYXRhbG9nCj4+CmVuZG9iago2IDAgb2JqCjw8Ci9BdXRob3IgKGFub255bW91cykg"
    "L0NyZWF0aW9uRGF0ZSAoRDoyMDI2MDkwMTAyNDEzMyswOCcwMCcpIC9DcmVhdG9yIChhbm9ueW1v"
    "dXMpIC9LZXl3b3JkcyAoKSAvTW9kRGF0ZSAoRDoyMDI2MDkwMTAyNDEzMyswOCcwMCcpIC9Qcm9k"
    "dWNlciAoUmVwb3J0TGFiIFBERiBMaWJyYXJ5IC0gXChvcGVuc291cmNlXCkpIAogIC9TdWJqZWN0"
    "ICh1bnNwZWNpZmllZCkgL1RpdGxlICh1bnRpdGxlZCkgL1RyYXBwZWQgL0ZhbHNlCj4+CmVuZG9i"
    "ago3IDAgb2JqCjw8Ci9Db3VudCAyIC9LaWRzIFsgMyAwIFIgNCAwIFIgXSAvVHlwZSAvUGFnZXMK"
    "Pj4KZW5kb2JqCjggMCBvYmoKPDwKL0ZpbHRlciBbIC9BU0NJSTg1RGVjb2RlIC9GbGF0ZURlY29k"
    "ZSBdIC9MZW5ndGggMTA1Cj4+CnN0cmVhbQpHYXBRaDBFPUYsMFVcSDNUXHBOWVReUUtrP3RjPklQ"
    "LDtXI1UxXjIzaWhQRU1fP0NXNEtJU2hbJnVHaj8raDdwdURPRXI0LG1qWipEQC1RKGwpQypZK14j"
    "LXE9czJvNj5RQCswJzs+fj5lbmRzdHJlYW0KZW5kb2JqCjkgMCBvYmoKPDwKL0ZpbHRlciBbIC9B"
    "U0NJSTg1RGVjb2RlIC9GbGF0ZURlY29kZSBdIC9MZW5ndGggMTA2Cj4+CnN0cmVhbQpHYXBRaDBF"
    "PUYsMFVcSDNUXHBOWVReUUtrP3RjPklQLDtXI1UxXjIzaWhQRU1fP0NXNEtJU2hbJnVHaj8raDdw"
    "dV81WVFQR3BzKCQvb0QiIk08ViIuJDZgT1orXi0jXSQ9IVI1VEcuI34+ZW5kc3RyZWFtCmVuZG9i"
    "agp4cmVmCjAgMTAKMDAwMDAwMDAwMCA2NTUzNSBmIAowMDAwMDAwMDYxIDAwMDAwIG4gCjAwMDAw"
    "MDAwOTIgMDAwMDAgbiAKMDAwMDAwMDE5OSAwMDAwMCBuIAowMDAwMDAwMzkyIDAwMDAwIG4gCjAw"
    "MDAwMDA1ODUgMDAwMDAgbiAKMDAwMDAwMDY1MyAwMDAwMCBuIAowMDAwMDAwOTE0IDAwMDAwIG4g"
    "CjAwMDAwMDA5NzkgMDAwMDAgbiAKMDAwMDAwMTE3NCAwMDAwMCBuIAp0cmFpbGVyCjw8Ci9JRCAK"
    "Wzw4YzVmZTk4YmY0ZmEyZjE2ZTFlZGU1YjgxOGNhN2RiZD48OGM1ZmU5OGJmNGZhMmYxNmUxZWRl"
    "NWI4MThjYTdkYmQ+XQolIFJlcG9ydExhYiBnZW5lcmF0ZWQgUERGIGRvY3VtZW50IC0tIGRpZ2Vz"
    "dCAob3BlbnNvdXJjZSkKCi9JbmZvIDYgMCBSCi9Sb290IDUgMCBSCi9TaXplIDEwCj4+CnN0YXJ0"
    "eHJlZgoxMzcwCiUlRU9GCg=="
)


def test_parse_document_supports_markdown_and_txt_and_normalizes(
    tmp_path: Path,
) -> None:
    workspace_root = tmp_path / "workspace"
    adapter = FileSystemAdapter(workspace_root)
    cases = [
        (
            Path("notes/README.MD"),
            b"\xef\xbb\xbfTitle\r\n\r\nBody\r",
            "markdown",
            "Title\n\nBody\n",
        ),
        (
            Path("notes/plain.TXT"),
            b"Plain\r\ntext",
            "text",
            "Plain\ntext",
        ),
    ]

    for file_entry_id, (
        relative_path,
        raw_bytes,
        source_format,
        expected_text,
    ) in enumerate(cases, start=7):
        file_path = workspace_root / relative_path
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(raw_bytes)

        document = load_document(
            adapter,
            workspace_id=3,
            file_entry_id=file_entry_id,
            source_relative_path=relative_path,
            document_id=DOCUMENT_ID,
        )

        assert document.source_relative_path == relative_path.as_posix()
        assert document.source_format == source_format
        assert document.normalized_text == expected_text
        assert document.source_version == sha256(raw_bytes).hexdigest()
        assert document.source_updated_at.tzinfo is not None


def test_parse_document_rejects_unsupported_format(tmp_path: Path) -> None:
    workspace_root = tmp_path / "workspace"
    unsupported_file = workspace_root / "notes" / "report.html"
    unsupported_file.parent.mkdir(parents=True)
    unsupported_file.write_bytes(b"not a supported text document")
    adapter = FileSystemAdapter(workspace_root)

    with pytest.raises(
        UnsupportedDocumentFormatError,
        match="unsupported document format",
    ):
        load_document(
            adapter,
            workspace_id=3,
            file_entry_id=7,
            source_relative_path=Path("notes/report.html"),
            document_id=DOCUMENT_ID,
        )


def test_load_document_supports_pdf_and_preserves_page_positions(
    tmp_path: Path,
) -> None:
    workspace_root = tmp_path / "workspace"
    pdf_file = workspace_root / "reports" / "summary.pdf"
    raw_bytes = base64.b64decode(PDF_FIXTURE_BASE64)
    pdf_file.parent.mkdir(parents=True, exist_ok=True)
    pdf_file.write_bytes(raw_bytes)
    adapter = FileSystemAdapter(workspace_root)

    document = load_document(
        adapter,
        workspace_id=3,
        file_entry_id=7,
        source_relative_path=Path("reports/summary.pdf"),
        document_id=DOCUMENT_ID,
    )

    assert document.source_format == "pdf"
    assert document.source_version == sha256(raw_bytes).hexdigest()
    assert [page.page_number for page in document.pages] == [1, 2]
    assert [
        document.normalized_text[page.start_offset : page.end_offset].strip()
        for page in document.pages
    ] == ["First PDF page", "Second PDF page"]
    assert document.pages[0].end_offset <= document.pages[1].start_offset


def test_load_document_rejects_malformed_pdf(tmp_path: Path) -> None:
    workspace_root = tmp_path / "workspace"
    malformed_file = workspace_root / "reports" / "broken.pdf"
    malformed_file.parent.mkdir(parents=True)
    malformed_file.write_bytes(b"%PDF-1.7\nnot a valid PDF")
    adapter = FileSystemAdapter(workspace_root)

    with pytest.raises(DocumentParseError, match="readable PDF"):
        load_document(
            adapter,
            workspace_id=3,
            file_entry_id=7,
            source_relative_path=Path("reports/broken.pdf"),
            document_id=DOCUMENT_ID,
        )


def test_load_document_supports_docx_and_preserves_source_positions(
    tmp_path: Path,
) -> None:
    workspace_root = tmp_path / "workspace"
    docx_file = workspace_root / "reports" / "summary.docx"
    docx_file.parent.mkdir(parents=True, exist_ok=True)

    word_document = WordDocument()
    word_document.add_paragraph("Intro paragraph")
    table = word_document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Left cell"
    table.cell(0, 1).text = "Right cell"
    word_document.add_paragraph("Closing paragraph")
    word_document.save(docx_file)
    raw_bytes = docx_file.read_bytes()
    adapter = FileSystemAdapter(workspace_root)

    document = load_document(
        adapter,
        workspace_id=3,
        file_entry_id=7,
        source_relative_path=Path("reports/summary.docx"),
        document_id=DOCUMENT_ID,
    )

    assert document.source_format == "docx"
    assert document.source_version == sha256(raw_bytes).hexdigest()
    assert document.pages == ()
    assert [position.element_type for position in document.source_positions] == [
        "paragraph",
        "table_cell",
        "table_cell",
        "paragraph",
    ]
    assert [
        document.normalized_text[position.start_offset : position.end_offset]
        for position in document.source_positions
    ] == [
        "Intro paragraph",
        "Left cell",
        "Right cell",
        "Closing paragraph",
    ]
    assert document.source_positions[0].paragraph_index == 0
    assert document.source_positions[1].table_index == 0
    assert document.source_positions[1].row_index == 0
    assert document.source_positions[1].cell_index == 0
    assert document.source_positions[2].cell_index == 1
    assert document.source_positions[3].paragraph_index == 1


def test_load_document_preserves_docx_heading_level(
    tmp_path: Path,
) -> None:
    workspace_root = tmp_path / "workspace"
    docx_file = workspace_root / "reports" / "heading.docx"
    docx_file.parent.mkdir(parents=True, exist_ok=True)

    word_document = WordDocument()
    word_document.add_heading("Overview", level=2)
    word_document.add_paragraph("Body")
    word_document.save(docx_file)
    adapter = FileSystemAdapter(workspace_root)

    document = load_document(
        adapter,
        workspace_id=3,
        file_entry_id=7,
        source_relative_path=Path("reports/heading.docx"),
        document_id=DOCUMENT_ID,
    )

    assert [position.heading_level for position in document.source_positions] == [
        2,
        None,
    ]


def test_load_document_preserves_docx_section_index(
    tmp_path: Path,
) -> None:
    workspace_root = tmp_path / "workspace"
    docx_file = workspace_root / "reports" / "sections.docx"
    docx_file.parent.mkdir(parents=True, exist_ok=True)

    word_document = WordDocument()
    word_document.add_paragraph("First section")
    word_document.add_section(WD_SECTION.NEW_PAGE)
    word_document.add_paragraph("Second section")
    word_document.save(docx_file)
    adapter = FileSystemAdapter(workspace_root)

    document = load_document(
        adapter,
        workspace_id=3,
        file_entry_id=7,
        source_relative_path=Path("reports/sections.docx"),
        document_id=DOCUMENT_ID,
    )

    nonempty_positions = [
        position
        for position in document.source_positions
        if position.start_offset < position.end_offset
    ]
    assert [
        (
            document.normalized_text[position.start_offset : position.end_offset],
            position.section_index,
        )
        for position in nonempty_positions
    ] == [("First section", 0), ("Second section", 1)]


def test_load_document_rejects_malformed_docx(tmp_path: Path) -> None:
    workspace_root = tmp_path / "workspace"
    malformed_file = workspace_root / "reports" / "broken.docx"
    malformed_file.parent.mkdir(parents=True)
    malformed_file.write_bytes(b"not a valid DOCX package")
    adapter = FileSystemAdapter(workspace_root)

    with pytest.raises(DocumentParseError, match="readable DOCX"):
        load_document(
            adapter,
            workspace_id=3,
            file_entry_id=7,
            source_relative_path=Path("reports/broken.docx"),
            document_id=DOCUMENT_ID,
        )
